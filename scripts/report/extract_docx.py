# -*- coding: utf-8 -*-
"""
Extract all content from a .docx file with proper UTF-8 output.
"""

import os
import sys
import io
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# ---------- paths ----------
DOCX_PATH = os.environ.get("DOCX_PATH", "./公园规划-设计.docx")  # Word文档路径
IMG_DIR   = os.environ.get("IMG_DIR", "./公园设计图")  # 图片输出目录
os.makedirs(IMG_DIR, exist_ok=True)

# ---------- open ----------
doc = Document(DOCX_PATH)

# =========================================================
# 1. TEXT PARAGRAPHS
# =========================================================
para_count = 0
lines = []
lines.append("=" * 70)
lines.append("TEXT CONTENT (paragraphs with styles)")
lines.append("=" * 70)
for i, para in enumerate(doc.paragraphs, start=1):
    text = para.text.strip()
    style = para.style.name if para.style else "None"
    if text or style.startswith("Heading"):
        para_count += 1
        if style.startswith("Heading"):
            level_str = style.replace("Heading", "").strip()
            prefix = "#" * int(level_str) if level_str.isdigit() else "##"
            lines.append(f"\n[P{i}] {style} | {prefix} {text}")
        else:
            lines.append(f"\n[P{i}] {style} | {text[:300]}")

# =========================================================
# 2. IMAGES
# =========================================================
img_count = 0
lines.append("\n" + "=" * 70)
lines.append("IMAGES")
lines.append("=" * 70)

image_parts = {}

for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_parts[rel.rId] = rel.target_part

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    drawing_elements = run._element.findall(qn('w:drawing'))
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall('.//' + qn('a:blip'))
                        for blip in blip_elements:
                            embed_id = blip.get(qn('r:embed'))
                            if embed_id and embed_id in doc.part.rels:
                                image_parts[embed_id] = doc.part.rels[embed_id].target_part

for idx, (rId, part) in enumerate(sorted(image_parts.items()), start=1):
    img_count += 1
    ext = os.path.splitext(part.partname)[-1]
    if not ext or ext == ".":
        ct = part.content_type
        ext_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/x-emf": ".emf",
            "image/x-wmf": ".wmf",
        }
        ext = ext_map.get(ct, ".png")
    fname = f"img_{img_count:03d}{ext}"
    fpath = os.path.join(IMG_DIR, fname)
    with open(fpath, "wb") as f:
        f.write(part.blob)
    lines.append(f"  [{idx}] {fname}  ({part.content_type}, {len(part.blob):,} bytes)")

lines.append(f"\n  Total images saved: {img_count}")

# =========================================================
# 3. TABLES
# =========================================================
table_count = 0
lines.append("\n" + "=" * 70)
lines.append("TABLES")
lines.append("=" * 70)
for t_idx, table in enumerate(doc.tables, start=1):
    table_count += 1
    rows = len(table.rows)
    cols = len(table.columns)
    lines.append(f"\n  Table {t_idx}: {rows} rows x {cols} cols")
    for r_idx, row in enumerate(table.rows, start=1):
        cell_texts = []
        for cell in row.cells:
            ct = cell.text.strip().replace("\n", " | ")
            cell_texts.append(ct)
        if r_idx <= 15:
            lines.append(f"    Row {r_idx}: {cell_texts}")
        elif r_idx == 16:
            lines.append(f"    ... ({rows - 15} more rows)")

lines.append(f"\n  Total tables: {table_count}")

# =========================================================
# SUMMARY
# =========================================================
lines.append("\n" + "=" * 70)
lines.append("SUMMARY")
lines.append("=" * 70)
lines.append(f"  Paragraphs with text/heading: {para_count}")
lines.append(f"  Images extracted:             {img_count}")
lines.append(f"  Tables:                       {table_count}")
lines.append(f"  Images saved to:              {IMG_DIR}")
lines.append("=" * 70)

output = "\n".join(lines)
print(output)

# Also write to a report file for easy reading
report_path = os.path.join(IMG_DIR, "_extraction_report.txt")
with open(report_path, "w", encoding="utf-8") as f:
    f.write(output)
print(f"\nReport saved to: {report_path}")
